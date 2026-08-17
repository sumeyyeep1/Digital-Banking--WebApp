from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:\Users\sümeyye\Downloads\klu_staj_defteri_şablon_-__temmuz_2026__ (1).docx")
OUTPUT_DIR = Path(r"C:\Users\sümeyye\Desktop\DigitalBankingSimple\outputs")
OUTPUT = OUTPUT_DIR / "staj_defteri_tekrarlar_duzeltildi.docx"
FALLBACK_OUTPUT = OUTPUT_DIR / "staj_defteri_tekrarlar_duzeltildi_v2.docx"


DAY_BODIES = {
    1: """Stajımın ilk gününde kuruma giriş süreci ve çalışacağım ortam hakkında genel bilgi aldım. Merkezi Sistemler ekibindeki sorumluluk alanlarını tanıyarak sunucu yönetimi, uygulama yayınlama, ağ yapılandırması ve veritabanı işlemlerinin kurum içindeki önemini inceledim. Ekipteki işlerin yalnızca teknik kurulumdan ibaret olmadığını, yapılan her değişikliğin çalışan uygulamaları ve kullanıcı erişimini etkileyebileceğini fark ettim.

Günün ilk bölümünde sanal makine kavramı üzerinde durdum. Fiziksel bir bilgisayar üzerinde birden fazla işletim sisteminin izole biçimde çalıştırılabildiğini ve bunun test, geliştirme, eğitim ve sunucu yönetimi için pratik bir yöntem olduğunu gördüm. Sanal makine oluştururken işlemci, bellek, disk ve ağ ayarlarının doğru seçilmesinin performans ve erişilebilirlik açısından önemli olduğunu değerlendirdim.

Daha sonra Windows Server kurulumu ve temel yapılandırma adımlarını takip ettim. Kurulumdan sonra sunucu adı, kullanıcı hesapları, IP ayarları ve güvenlik seçenekleri gibi başlangıç ayarlarını inceledim. Server Manager ekranı üzerinden rollerin ve servislerin yönetilebildiğini gördüm. Bu çalışma, ilerleyen günlerde IIS, veritabanı ve uzak bağlantı işlemlerini daha rahat anlamam için temel oluşturdu.

Ayrıca kurum içinde kullanılan Excel e-posta gönderim aracı hakkında bilgi aldım. Excel dosyasındaki belirli verilerin okunarak ilgili kişilere otomatik e-posta gönderilmesi mantığını inceledim. Bu örnek sayesinde bir iş ihtiyacının küçük bir yazılım aracıyla nasıl otomatikleştirilebildiğini ve düzenli yapılan işlemlerde zaman kazandırabildiğini gördüm. İlk gün, sistem tarafındaki temel kavramları tanımam ve sonraki çalışmalara hazırlanacak ortamı anlamam açısından verimli geçti.""",
    2: """İkinci gün Windows Server mimarisi ve uzak yönetim konusu üzerinde çalıştım. Server Manager ekranındaki Dashboard, Local Server, All Servers ve Add Roles and Features bölümlerini inceleyerek sunucuların merkezi bir yerden nasıl yönetildiğini kavradım. Bir sunucunun yalnızca işletim sistemi kurulu bir makine olmadığını, üzerinde çalışan roller ve servisler sayesinde kurumsal uygulamalara hizmet verdiğini gördüm.

Rol ve özellik kavramlarını karşılaştırdım. IIS gibi rollerin sunucuya belirli bir görev kazandırdığını, ek özelliklerin ise bu görevleri destekleyen yardımcı bileşenler olduğunu inceledim. Kurulum sihirbazında seçimlerin bilinçli yapılması gerektiğini, gereksiz servislerin güvenlik ve performans açısından risk oluşturabileceğini değerlendirdim.

Günün önemli kısmını RDP ile uzaktan bağlantı konusuna ayırdım. Remote Desktop Protocol kullanılarak farklı bir bilgisayardan Windows Server arayüzüne erişilebildiğini uygulamalı olarak gördüm. Bağlantı için IP adresi, kullanıcı bilgileri, parola ve güvenlik duvarı kurallarının uyumlu olması gerektiğini test ettim. Yanlış kullanıcı yetkisi veya kapalı firewall kuralı olduğunda bağlantının kurulamadığını gözlemledim.

Uzak yönetim sırasında güvenlik konusunun özellikle dikkat gerektirdiğini fark ettim. Her kullanıcıya yönetici yetkisi verilmemesi, güçlü parola kullanılması ve erişimlerin kontrol altında tutulması gerektiği anlatıldı. Bu çalışmalar sonucunda Windows Server üzerinde rol ekleme, temel servisleri takip etme ve RDP ile güvenli bağlantı kurma sürecini daha net hale getirdim.""",
    3: """Üçüncü gün ASP.NET tabanlı bir uygulamayı kendi local laptopumda yayınlama ve çalıştırma sürecine odaklandım. Şirket ortamında doğrudan yetkim olmadığı için sunucu üzerindeki adımları yetkili kişilerden dinleyerek takip ettim; uygulama tarafındaki denemeleri ise kendi local ortamımda yaptım. Böylece yazılan kodun geliştirme ortamından çalışır bir web uygulamasına nasıl dönüştüğünü daha somut biçimde gördüm.

Local ortamda publish çıktısı almayı, uygulama dosyalarının hangi klasöre üretildiğini ve tarayıcıdan uygulamaya nasıl erişildiğini test ettim. IIS tarafında ise web sitesi oluşturma, application pool seçme, physical path belirleme ve binding ayarlarının şirket ortamında nasıl kullanıldığını gözlemledim. Application pool yapısının uygulamanın çalışma ortamını yönettiğini, farklı uygulamaların birbirinden ayrılmasını sağladığını ve hata durumlarında yönetimi kolaylaştırdığını değerlendirdim.

Yayınlama sırasında klasör izinleri üzerinde durdum. Kendi bilgisayarımda dosya yolu ve çalışma klasörü ayarlarını kontrol ettim; şirket sunucularında ise izinlerin yetkili kişiler tarafından yönetildiğini gördüm. Uygulama dosyalarının bulunduğu dizine ilgili çalışma kullanıcısının erişebilmesi gerektiğini, aksi halde dosya okuma veya yazma hataları oluşabileceğini not ettim.

Son bölümde localde yayınladığım uygulamayı tarayıcıdan açarak sonucu kontrol ettim. Hata alındığında uygulama logları, IIS kullanılan ortamlarda IIS logları ve Windows Event Viewer kayıtlarının incelenebileceğini öğrendim. Bu gün, kendi bilgisayarımda çalışan bir ASP.NET uygulamasının yayınlama mantığını uygulayarak kavramamı sağladı.""",
    4: """Dördüncü gün daha çok hata giderme ve sanal ağ yapılandırmaları üzerine çalıştım. Bir uygulama sunucu üzerinde yayınlandığında karşılaşılan sorunların yalnızca koddan kaynaklanmadığını; port, firewall, izin, bağlantı dizesi, uygulama havuzu ve ağ ayarlarının da sonucu etkilediğini inceledim. Bu nedenle hata çözümünde tek bir noktaya bakmak yerine süreci adım adım kontrol etmek gerektiğini gördüm.

IIS üzerinde oluşabilecek yaygın hataları değerlendirdim. 404 durumunda dosya yolu veya route ayarlarının, 500 durumunda uygulama içi hataların, 401 ve 403 durumlarında ise yetkilendirme ve erişim izinlerinin kontrol edilmesi gerektiğini karşılaştırdım. Event Viewer ve IIS logları üzerinden hata zamanını, istek adresini ve dönen durum kodunu takip etmeyi denedim.

Sanal makine ağ ayarlarında NAT ve Bridged seçeneklerini inceledim. NAT modunda sanal makinenin fiziksel bilgisayarın ağı üzerinden dışarıya çıktığını, Bridged modda ise ağda ayrı bir cihaz gibi davranabildiğini gördüm. Bu farkın özellikle sunucuya dışarıdan erişim testlerinde önemli olduğunu değerlendirdim. IP adresi değiştiğinde bağlantıların etkilenebileceğini ve statik IP kullanımının sunucu ortamlarında daha düzenli bir yapı sağladığını fark ettim.

Çalışma sonunda bir uygulamanın çalışmamasına neden olan sorunları sınıflandırmayı denedim: uygulama hatası, sunucu yapılandırması, ağ erişimi veya yetki problemi. Bu ayrım, sonraki günlerde deployment ve veritabanı bağlantısı kontrollerini daha sistemli yapmam için yol gösterici oldu.""",
    5: """Beşinci gün ASP.NET uygulama geliştirme, local deployment ve veritabanı bağlantısı konularını birlikte ele aldım. Şirket sunucularında işlem yetkim olmadığı için canlı ortamdaki adımları gözlemledim; kendi laptopumda ise uygulamayı yayınlayıp bağlantı ayarlarını test ettim. Böylece geliştirme, yayınlama ve test adımlarının birbirinden ayrı fakat bağlantılı süreçler olduğunu gördüm.

Deployment sırasında uygulama dosyalarının doğru klasöre üretilmesini, localde doğru adresten açılmasını ve hedef framework sürümüyle uyumlu çalışmasını kontrol ettim. IIS kullanılan şirket ortamlarında ise site klasörü ve application pool ayarlarının nasıl önem kazandığını gözlemledim. Yayınlama sonrasında yalnızca ana sayfanın açılmasını yeterli görmeyip farklı sayfa ve işlem akışlarını da denedim.

Veritabanı tarafında SQL Server bağlantı mantığını kendi local ayarlarım üzerinden tekrar ettim. Connection string içinde sunucu adı, veritabanı adı, kimlik doğrulama yöntemi ve gerekli güvenlik ayarlarının bulunduğunu kontrol ettim. Windows Authentication ve SQL Authentication arasındaki farkı karşılaştırdım. Yanlış bağlantı dizesi, eksik kullanıcı yetkisi veya kapalı servis gibi durumların uygulamanın veri tabanına erişmesini engelleyebileceğini gördüm.

Ayrıca yedek alma ve geri yükleme işlemlerinin neden gerekli olduğunu değerlendirdim. Veritabanı değişikliklerinden önce yedek alınmasının olası veri kayıplarına karşı güvence sağladığını fark ettim. Bu günün çalışması, uygulama yayınlama ile veritabanı erişiminin birlikte düşünülmesi gerektiğini ve canlı ortama çıkmadan önce bağlantıların dikkatli test edilmesinin önemini pekiştirdi.""",
    6: """Altıncı gün SQL Server üzerinde veritabanı oluşturma ve sunucuları canlı ortama hazırlama süreçlerine odaklandım. Şirket ortamında bu işlemleri yapma yetkim olmadığı için Veritabanı Yöneticisi ile SQL Server Management Studio üzerindeki adımları gözlemledim; kendi localimde ise aynı kavramları not alarak pekiştirdim. Veri dosyası, log dosyası, başlangıç boyutu ve büyüme ayarlarını inceledim. Verilerin .mdf, işlem kayıtlarının ise .ldf dosyalarında tutulduğunu görerek veritabanının arka plandaki dosya düzenini daha iyi kavradım.

Kullanıcı ve yetkilendirme tarafında login ile database user arasındaki ilişkiyi inceledim. Bir kullanıcının sunucuya bağlanabilmesi ile belirli bir veritabanında işlem yapabilmesinin farklı yetkilere bağlı olduğunu gördüm. Kullanıcılara yalnızca ihtiyaç duydukları rollerin verilmesi gerektiği özellikle vurgulandı. Bu yaklaşımın hem veri güvenliği hem de hatalı işlem riskini azaltmak için önemli olduğunu değerlendirdim.

Günün devamında Sistem Yöneticisi ile dört farklı Linux sunucunun canlı ortama alınması sürecini gözlemledim. Sunucuların IP adresleri, host adları, SSH erişimleri, servis durumları ve sistem güncellemeleri yetkili kişiler tarafından kontrol edildi. Production ortamında yapılan işlemlerin doğrudan çalışan hizmetleri etkileyebileceğini gördüğüm için değişikliklerin planlı, kontrollü ve test edilerek uygulanması gerektiğini not ettim.

SSH bağlantısı üzerinden temel sistem kontrollerinin nasıl yapıldığını izledim. nmtui aracıyla statik IP, ağ geçidi ve DNS ayarlarının düzenlenebildiğini inceledim. Yanlış ağ yapılandırmasının sunucunun erişilebilirliğini tamamen etkileyebileceğini gördüm. yum update ile paket güncelleme sürecinde bağımlılıkların kontrol edilmesi ve güncelleme sonrası kritik servislerin tekrar denenmesi gerektiği belirtildi. Son olarak SIEM kavramını tanıyarak sunucu ve uygulama loglarının merkezi olarak izlenmesinin güvenlik takibi açısından önemini değerlendirdim.""",
    7: """Yedinci gün PostgreSQL veritabanının Docker container üzerinde çalıştırılması ve temel Linux komutları üzerine çalıştım. Şirket sistemlerinde yetkim olmadığı için üretim ortamında müdahale yapmadım; Docker ve PostgreSQL denemelerini kendi localimde gerçekleştirerek sürecin mantığını kavramaya çalıştım. Daha önce klasik kurulumla çalışan veritabanı yaklaşımını, Docker ile izole çalışan bir servis haline getirme fikriyle karşılaştırdım.

Docker tarafında image ve container kavramlarını ayırdım. Image yapısının çalıştırılacak ortamın hazır şablonu, container yapısının ise bu şablonun çalışan örneği olduğunu inceledim. PostgreSQL image kullanılarak container oluşturulurken veritabanı adı, kullanıcı adı, parola ve port bilgileri gibi ortam değişkenlerinin tanımlanması gerektiğini gördüm. Port mapping sayesinde container içindeki PostgreSQL servisinin local makineden erişilebilir hale getirildiğini test ettim.

Verilerin container silindiğinde kaybolmaması için volume yapısını kullandım. PostgreSQL verilerinin kalıcı bir disk alanında tutulması, container yeniden başlatılsa bile verinin korunmasını sağladı. docker ps ile çalışan containerları listelemeyi, docker logs ile servis kayıtlarını incelemeyi, docker start ve docker stop ile container durumunu yönetmeyi denedim. Bu komutlar hata takibi ve servis kontrolü için pratik bir bakış kazandırdı.

Günün ikinci bölümünde Linux komut satırında dosya, klasör, servis ve ağ kontrolleri yaptım. pwd, ls, cd, mkdir, cp, mv ve rm komutlarıyla dizin işlemlerini; cat, less ve grep ile dosya içeriği incelemeyi kullandım. top, df -h, free -m, ip addr, ping, ss, chmod, chown ve systemctl komutlarının sistem yönetimindeki yerini gördüm. Böylece hem container tabanlı veritabanı yönetimini hem de Linux sunucularda temel kontrol adımlarını uygulamış oldum.""",
    8: """Sekizinci gün Merkezi Sistemler ekibindeki çalışmalarımı tamamlayarak Yazılım Geliştirme ekibine geçtim. Bu geçişle birlikte çalışma konum sistem yönetiminden yazılım geliştirme sürecine kaydı. Ekipteki görev dağılımını, backend ve frontend taraflarının nasıl ayrıldığını, mobil geliştirme ile API geliştirme arasındaki ilişkiyi dinledim. Böylece bir yazılım projesinde farklı uzmanlıkların aynı hedef için nasıl birlikte çalıştığını daha iyi gördüm.

İlk olarak Git ve sürüm kontrol mantığı üzerinde çalıştım. Commit, branch, checkout ve merge işlemlerini Learn Git Branching üzerinde uyguladım. Branch kullanımının yeni özellikleri ana kodu bozmadan geliştirmeyi sağladığını, commit geçmişinin ise yapılan değişiklikleri izlenebilir hale getirdiğini değerlendirdim. Rebase, reset, revert ve cherry-pick kavramlarını temel seviyede tanıdım; bu işlemlerin dikkatli kullanılmadığında proje geçmişini etkileyebileceğini not ettim.

Günün devamında staj süresince geliştireceğim proje için fikir değerlendirmesi yaptık. Digital Banking projesinin kullanıcı kaydı, giriş, hesap yönetimi, para yatırma, para çekme, transfer, kart işlemleri ve arayüz gibi temel bölümlerden oluşabileceği konuşuldu. Önceliğin kısa sürede çalışan bir temel sürüm hazırlamak, ardından zaman kalırsa ek özellikler ve iyileştirmeler eklemek olması kararlaştırıldı.

Agile yaklaşımındaki sprint, backlog, user story ve daily meeting kavramlarını da inceledim. Projeyi küçük parçalara bölerek ilerlemenin hem planlamayı kolaylaştırdığını hem de düzenli geri bildirim almayı sağladığını gördüm. Bu gün, Digital Banking projesi için yol haritasını belirlediğim ve yazılım geliştirme tarafındaki çalışma düzenine uyum sağlamaya başladığım aşama oldu.""",
    9: """Dokuzuncu gün kendi localimde geliştireceğim Digital Banking Web API projesinin temel mimari planlamasına başladım. İlk tasarımda N-Katmanlı Clean Architecture yaklaşımını esas aldım. Domain, Application, Infrastructure ve API katmanlarının hangi sorumlulukları taşıyabileceğini inceleyerek kodu yalnızca çalışan bir yapı olarak değil, sürdürülebilir ve anlaşılır bir proje olarak kurmanın önemini değerlendirdim.

Domain tarafında banka uygulamasında ihtiyaç duyulabilecek temel varlıkları belirledim. User, Customer, Account ve Transaction gibi sınıfların hangi bilgileri tutacağını düşündüm. Tüm varlıklarda ortak kullanılabilecek Id, oluşturulma tarihi, güncellenme tarihi ve kayıt durumu gibi alanlar için BaseEntity fikrini çalıştım. AccountType, Currency, TransactionType, UserRole ve EntityStatus gibi enum yapılarını kullanmanın sabit değerleri daha kontrollü yönetmeye yardımcı olduğunu gördüm.

Veritabanı yaklaşımı olarak Entity Framework Core Code-First mantığını inceledim. Model sınıflarından veritabanı şeması üretilmesi, migration dosyalarıyla değişikliklerin takip edilmesi ve AppDbContext üzerinden tablolarla ilişki kurulması konularını değerlendirdim. Parasal alanlarda decimal kullanılmasının, tarih alanlarında ise tutarlı bir zaman formatı seçilmesinin bankacılık benzeri işlemler için önemli olduğunu not ettim.

Bu aşamada proje dosyalarının ilk planı ve teknik kararlar üzerinde durdum. Git geçmişinde Digital Banking projesine ait ilk kayıtlar 4 Ağustos tarihinde başladığı için bu gündeki çalışmayı kesin bir commit çıktısı olarak değil, proje hazırlığı ve mimari tasarım aşaması olarak ele aldım. Çalışmanın sonucu olarak sonraki günlerde uygulanacak entity, DTO, controller, service ve veritabanı yapısı için daha net bir taslak oluştu.""",
    10: """Onuncu gün kimlik doğrulama ve API test süreci üzerine çalıştım. Digital Banking projesinde kullanıcı kaydı ve giriş işlemlerinin nasıl kurgulanması gerektiğini inceledim. Kullanıcının sisteme girebilmesi için e-posta ve parola bilgisinin alınması, kayıt sırasında tekrar eden kullanıcıların kontrol edilmesi ve giriş başarılı olduğunda istemciye bir yanıt dönülmesi gerektiğini belirledim.

DTO kullanımını bu akış üzerinden ele aldım. LoginRequestDto ve RegisterRequestDto gibi sınıfların istemciden gelen verileri taşıması, LoginResponseDto yapısının ise kullanıcıya dönecek token ve temel kullanıcı bilgilerini içermesi planlandı. Entity ile DTO arasındaki farkı bu örnek üzerinde daha net ayırdım: Entity veritabanı tablosunu temsil ederken DTO, API üzerinden taşınacak veri şeklini belirliyordu.

Swagger arayüzünü kullanarak endpointlerin nasıl test edilebileceğini inceledim. Bir API metoduna hangi HTTP yöntemiyle istek gönderileceği, Body bölümüne JSON verisinin nasıl yazılacağı ve dönen HTTP durum kodlarının nasıl yorumlanacağı üzerinde durdum. 200, 400, 401 ve 404 gibi durum kodlarının kullanıcıya veya geliştiriciye işlemin sonucu hakkında bilgi verdiğini gördüm.

Bu çalışmalar sırasında kimlik doğrulama sürecinin yalnızca giriş ekranından ibaret olmadığını fark ettim. Şifre saklama, token üretimi, endpoint koruması ve kullanıcı bilgisinin sonraki isteklerde güvenli biçimde taşınması birlikte düşünülmesi gereken konulardı. Günün çıktısı, AuthController ve AuthService tarafında geliştirilecek akışın daha anlaşılır hale gelmesi oldu.""",
    11: """On birinci gün C# dilindeki temel yapıları, Dependency Injection ve asenkron programlama konularını proje ihtiyaçlarıyla ilişkilendirerek çalıştım. Class, property, get/set, constructor, abstract class, inheritance, virtual ve override kavramlarını tekrar ettim. Bu yapıları yalnızca teorik olarak değil, Digital Banking projesindeki model ve service sınıflarını anlamak için gerekli temel bilgiler olarak değerlendirdim.

Constructor yapısının bir sınıf oluşturulurken ihtiyaç duyduğu bağımlılıkları almasını sağladığını inceledim. Dependency Injection sayesinde controller veya service sınıflarının ihtiyaç duyduğu AppDbContext, IAuthService ya da IAccountService gibi nesnelerin Program.cs içerisinde tanımlanıp çalışma zamanında otomatik verildiğini gördüm. Böylece sınıflar arasında daha düzenli ve test edilebilir bir bağ kurulabildiğini kavradım.

LINQ sorgularını koleksiyonlar ve Entity Framework örnekleri üzerinden çalıştım. Where ile filtreleme, Select ile veri dönüştürme, AnyAsync ile kayıt var mı kontrolü, FirstOrDefaultAsync ile uygun ilk kaydı alma ve ToListAsync ile listeleme işlemlerinin projede sık kullanılacağını gördüm. Bu metotların özellikle kullanıcı, hesap ve işlem kayıtlarını sorgularken controller yerine service katmanında yer almasının daha doğru olduğunu değerlendirdim.

Asenkron programlama tarafında Task, Task<T>, async ve await kavramlarını inceledim. Veritabanı ve dış API isteklerinde işlemin tamamlanmasını beklerken uygulamanın kaynakları daha verimli kullanabilmesi için async yapının tercih edildiğini gördüm. try-catch bloklarıyla hata yönetimi konusunu da çalışarak, beklenmeyen durumlarda kullanıcıya kontrollü yanıt verilmesinin önemini pekiştirdim.""",
    12: """On ikinci gün ASP.NET Core Web API yapısını controller ve service ilişkisi üzerinden detaylandırdım. [ApiController], [Route], [HttpGet], [HttpPost] ve [HttpPut] attribute'larının endpointlerin adresini ve hangi HTTP metoduyla çalışacağını belirlediğini inceledim. [Authorize] attribute'unun ise yalnızca kimliği doğrulanmış kullanıcıların ilgili endpointlere erişmesini sağladığını gördüm.

İstemciden veri alma yöntemlerini karşılaştırdım. [FromBody] ile JSON verisinin DTO nesnesine dönüştürülmesini, [FromRoute] ile URL içindeki değerlerin alınmasını ve [FromQuery] ile sorgu parametrelerinin metoda aktarılmasını inceledim. Request DTO ve Response DTO ayrımı, API'nin hem daha kontrollü hem de daha güvenli veri taşımasını sağladı.

IActionResult ve ActionResult<T> dönüş tiplerinin API yanıtlarını standartlaştırmak için kullanıldığını gördüm. Ok(), BadRequest(), Unauthorized() ve NotFound() gibi yanıtların farklı işlem sonuçlarını HTTP durum kodlarıyla ifade ettiğini test ettim. Bu yaklaşım frontend tarafının gelen cevaba göre kullanıcıya doğru mesaj göstermesini kolaylaştırıyordu.

Entity Framework Core tarafında AppDbContext'in Dependency Injection ile servislere verildiğini, _context.Accounts gibi DbSet property'lerinin veritabanındaki tablolarla eşleştiğini inceledim. Service katmanında yer alan metotlar controllerdan gelen isteği alıyor, iş kuralını uyguluyor ve gerekirse AppDbContext üzerinden veritabanına erişiyordu. Bu gün sonunda HTTP isteğinin Controller -> Service -> AppDbContext -> Veritabanı şeklindeki temel akışı benim için daha anlaşılır hale geldi.""",
    13: """On üçüncü gün projenin mimari düzenini sadeleştirme ve dosyaları mevcut ihtiyaçlara göre yeniden konumlandırma çalışması yaptım. Başlangıçta planlanan çok katmanlı yapı, staj süresi ve projenin kapsamı dikkate alınarak Controllers, Services, Interfaces, DTOs, Models ve Data klasörlerinden oluşan daha anlaşılır bir yapıya dönüştürüldü. Bu değişikliği bir refactoring çalışması olarak ele aldım.

Controllers klasöründe HTTP isteklerini karşılayan sınıflar, Services klasöründe iş kurallarını ve veritabanı işlemlerini yürüten sınıflar, Interfaces klasöründe ise servislerin metot sözleşmeleri yer aldı. Controllerların doğrudan somut sınıfa değil interface yapısına bağımlı olması, kodun daha gevşek bağlı ilerlemesini sağladı. Bu ayrım sayesinde örneğin AccountsController yalnızca isteği alıp IAccountService metodunu çağırıyor, hesap oluşturma veya listeleme detayları AccountService içinde kalıyordu.

DTOs klasöründe Auth, Accounts, Transactions ve Cards altında request ve response sınıfları düzenlendi. Models klasöründe User, Customer, Account, Transaction, Card ve BaseEntity sınıfları yer aldı. Enums klasöründe UserRole, AccountType, Currency, CardType, TransactionType ve EntityStatus sabit değerleri toplandı. Bu düzenleme, proje büyüdükçe dosyaların daha kolay bulunmasını ve her sınıfın görevini daha anlaşılır kılmasını sağladı.

Data klasöründeki AppDbContext sınıfında DbSet tanımları, ilişki ayarları, benzersiz indeksler ve decimal hassasiyetleri incelendi. Migration dosyalarının model değişikliklerini veritabanına yansıtmak için oluşturulduğunu gördüm. Mimari sadeleşince proje yapısı sunumda anlatılabilecek kadar netleşti ve sonraki günlerde auth, hesap, kart ve işlem servislerini bu düzen içinde geliştirmek kolaylaştı.""",
    14: """On dördüncü gün kullanıcı kaydı, giriş işlemi ve JWT tabanlı yetkilendirme üzerinde çalıştım. AuthController içinde /api/auth/register ve /api/auth/login endpointleri, AuthService tarafındaki metotlara yönlendirildi. Böylece controller yalnızca isteği karşılayan katman olarak kaldı, asıl iş kuralları service sınıfında toplandı.

Kayıt işleminde RegisterRequestDto ile gelen e-posta, parola ve kullanıcı bilgileri alındı. Yeni kullanıcı oluşturulmadan önce aynı e-posta veya TC kimlik numarasıyla kayıt olup olmadığı Entity Framework sorgularıyla kontrol edildi. Uygun durumda User ve Customer kayıtları oluşturularak AppDbContext üzerinden veritabanına eklendi. Bu akış, kullanıcı kaydı sırasında hem veri tekrarını engellemeyi hem de müşteri bilgisiyle kullanıcı bilgisini ilişkilendirmeyi sağladı.

Şifre güvenliği için PasswordHasher kullandım. Parolanın veritabanına açık metin olarak kaydedilmemesi gerektiğini, hash değerinin tek yönlü bir temsil sunduğunu uygulama içinde gördüm. Giriş sırasında LoginRequestDto ile gelen parola, kayıtlı hash değerle karşılaştırıldı. Doğrulama başarılı olduğunda LoginResponseDto içinde JWT token ve temel kullanıcı bilgileri döndürüldü.

JWT token içine kullanıcının kimliği, e-posta adresi ve rolü claim olarak eklendi. Program.cs dosyasında JWT Bearer ayarları, token imzası, geçerlilik süresi ve authentication middleware sırası yapılandırıldı. [Authorize] eklenen endpointlerde token gönderilmediğinde 401 Unauthorized cevabı alındığını, geçerli token ile isteğin işleme alındığını test ettim. Bu çalışma, güvenli API erişiminin kayıt ve girişten sonraki tüm işlemler için temel olduğunu gösterdi.""",
    15: """On beşinci gün Digital Banking projesinin temel bankacılık işlemlerini geliştirdim. AccountsController ve AccountService üzerinden kullanıcının kendi hesaplarını listeleme, yeni hesap oluşturma ve hesap bilgilerini güncelleme akışlarını inceledim. Kullanıcı kimliği JWT içindeki NameIdentifier claim'inden alınarak her hesap sorgusunda sahiplik kontrolü yapıldı. Bu sayede bir kullanıcının başka bir kullanıcının hesabını görüntülemesi veya değiştirmesi engellendi.

Yeni hesap oluştururken CreateAccountRequestDto ile hesap türü ve para birimi bilgileri alındı. Account entity'si oluşturulurken benzersiz IBAN üretimi yapıldı ve aynı IBAN'ın daha önce veritabanında bulunup bulunmadığı kontrol edildi. Hesap güncelleme işleminde ise UpdateAccountRequestDto ile yalnızca izin verilen alanların değiştirilmesi sağlandı. Bu süreç, DTO kullanımının veri sınırlarını belirlemek açısından neden gerekli olduğunu gösterdi.

Kart işlemlerinde CardsController, ICardService ve CardService yapısını kullandım. Kullanıcının kartlarını listeleme, yeni kart oluşturma ve kart bilgilerini güncelleme akışlarını geliştirdim. Kart oluşturulmadan önce kartın bağlanacağı hesabın kullanıcıya ait olup olmadığı kontrol edildi. API yanıtında kart numarasının yalnızca son dört hanesi gösterildi, CVV bilgisi ise açık metin olarak tutulmadı. Böylece hassas kart verisinin response DTO ile dışarıya kontrollü şekilde aktarılması sağlandı.

Günün devamında TransactionsController ve TransactionService üzerinden para yatırma, para çekme ve transfer işlemlerini tamamladım. Para çekmede bakiye yeterliliği, transferde gönderen hesabın kullanıcıya ait olması, alıcı hesabın bulunması, hesapların farklı olması ve para birimlerinin uyumu kontrol edildi. Başarılı transferde gönderen bakiyesi azaltılıp alıcı bakiyesi artırıldı ve hareket kaydı Transaction tablosuna yazıldı. Bu aşama ile projenin temel bankacılık işlevleri çalışır duruma geldi.""",
    16: """On altıncı gün kendi localimde çalışan React frontend arayüzünü oluşturmaya ve backend ile iletişimini kurmaya odaklandım. Vite ile hazırlanan frontend tarafında sayfa yapısını, component kullanımını ve merkezi API servis dosyasını inceledim. LoginPage, DashboardPage, AccountsPage, TransferPage, TransactionsPage ve CardsPage gibi sayfalar üzerinden kullanıcının uygulamadaki temel akışlarını tasarladım.

frontend/src/services/api.ts dosyasının backend endpointlerine istek gönderen merkezi yapı olduğunu gördüm. Bu dosyada API base URL bilgisi kullanılarak login, register, hesap listeleme, kart işlemleri, transfer ve piyasa verisi istekleri tanımlandı. Kullanıcı giriş yaptıktan sonra alınan JWT token'ın sonraki isteklerde Authorization header içine Bearer token olarak eklenmesi sağlandı. Böylece [Authorize] ile korunan backend endpointlerine frontend üzerinden erişim kurulabildi.

useAuth hook'u ile oturum bilgisinin frontend tarafında yönetilmesini inceledim. Kullanıcı bilgisi ve token sessionStorage içinde tutuldu, sayfa yenilendiğinde oturumun devam etmesi sağlandı. Bu yapının kullanıcı deneyimini kolaylaştırdığını ancak hassas bilgilerin frontendde saklanırken dikkatli olunması gerektiğini değerlendirdim.

Backend ile frontend arasındaki iletişimde CORS ayarlarının önemini test ettim. React uygulaması farklı porttan çalıştığı için Program.cs içinde izin verilen originler tanımlandı. API yanıtlarına göre loading, error, boş liste ve başarılı işlem durumları arayüzde gösterildi. Bu gün sonunda backendde çalışan endpointlerin kullanıcı tarafından görülebilen ekranlara bağlanmasıyla proje daha kullanılabilir bir hale geldi.""",
    17: """On yedinci gün kullanıcı deneyimi, backend uyumluluğu ve Docker yapılandırması üzerinde çalıştım. Frontend sayfalarında kullanıcıya gösterilen metinleri, hata durumlarını ve işlem sonrası geri bildirimleri gözden geçirdim. API'den dönen cevapların doğrudan kullanıcıya anlaşılır biçimde yansıtılması, uygulamanın sadece teknik olarak değil kullanım açısından da düzgün çalışması için gerekliydi.

Backend tarafında mevcut kayıtlar ve şifre doğrulama akışı incelendi. Daha önce düz metin olarak kalmış parola değerleriyle hashlenmiş parola değerleri arasında uyumluluk ihtiyacı değerlendirildi. AuthService içinde giriş sırasında parolanın güvenli doğrulanması ve gerekli durumda eski kayıtların daha güvenli hale getirilmesi üzerinde çalışıldı. Bu konu, güvenlik iyileştirmelerinin bazen mevcut verilerle uyumu da dikkate alması gerektiğini gösterdi.

Docker yapılandırmasında backend, frontend ve veritabanı servislerinin ayrı containerlar olarak nasıl çalışabileceğini inceledim. Backend Dockerfile .NET uygulamasını build edip runtime ortamında çalıştıracak şekilde, frontend Dockerfile ise React uygulamasını build edip Nginx ile yayınlayacak şekilde düzenlendi. docker-compose.yml dosyasında servislerin birlikte ayağa kalkması, port eşleşmeleri ve environment variable kullanımı ele alındı.

Bu gün ayrıca .gitignore ve ortam ayarlarının önemini değerlendirdim. Bağlantı bilgileri, API anahtarları ve JWT secret gibi hassas değerlerin kaynak koduna açık biçimde eklenmemesi gerektiğini gördüm. Docker ve frontend-backend uyumluluğu sayesinde proje farklı ortamlarda daha kolay çalıştırılabilecek bir yapıya yaklaştı.""",
    18: """On sekizinci gün Entity Framework Core ilişkileri, migration yapısı ve AppDbContext üzerinde daha ayrıntılı çalıştım. AppDbContext dosyasında User, Customer, Account, Transaction ve Card entity'lerinin DbSet olarak tanımlandığını inceledim. Bu DbSet'lerin C# tarafındaki sınıfları veritabanındaki tablolarla eşleştirdiğini ve sorguların bu yapı üzerinden yazıldığını gördüm.

Model ilişkilerini primary key ve foreign key mantığıyla değerlendirdim. Kullanıcı ile müşteri, müşteri ile hesap, hesap ile kart ve hesaplar ile transaction kayıtları arasındaki bağlantıların veri bütünlüğü açısından gerekli olduğunu kavradım. DeleteBehavior.Restrict gibi ayarların, hesap silme durumunda geçmiş işlem kayıtlarının yanlışlıkla etkilenmesini önlemek için kullanıldığını inceledim.

Fluent API ile yapılan yapılandırmaları karşılaştırdım. Email, IdentityNumber, Iban ve CardNumber gibi alanlarda benzersiz indeks kullanılması, aynı değerin birden fazla kayıtta oluşmasını engelliyordu. Decimal alanlarda precision tanımlanması, bakiye ve işlem tutarı gibi parasal değerlerin doğru saklanması için önemliydi. Bu ayarlar, bankacılık benzeri bir uygulamada veritabanı şemasının dikkatli hazırlanması gerektiğini gösterdi.

Migration dosyalarının model değişikliklerinin geçmişini tuttuğunu ve database update ile veritabanına uygulandığını inceledim. Include, Where, Select, AnyAsync, FirstOrDefaultAsync, ToListAsync ve SaveChangesAsync metotlarının service sınıflarında nasıl kullanıldığını değerlendirdim. Bu çalışma sonucunda veritabanı işlemlerinin controller içinde değil, service katmanı ve AppDbContext üzerinden yönetildiğini daha net şekilde açıklayabilecek hale geldim.""",
    19: """On dokuzuncu gün dış piyasa verisi entegrasyonu ve kart yönetimi arayüzü üzerinde çalıştım. CollectAPI üzerinden altın, döviz, borsa ve hisse senedi verilerinin alınması için MarketController endpointlerini ve IMarketService/MarketService yapısını inceledim. Bu özellik sayesinde uygulama yalnızca kendi veritabanındaki bankacılık işlemlerini değil, dış bir servisten gelen güncel piyasa bilgilerini de gösterebilir hale geldi.

MarketService içinde BaseUrl ve ApiKey bilgileri appsettings veya environment variable üzerinden okunacak şekilde düzenlendi. API anahtarının koda açık biçimde yazılmaması gerektiğini, geliştirme ortamında appsettings.Development.json veya ortam değişkeniyle yönetilmesinin daha doğru olduğunu değerlendirdim. CollectAPI isteklerinde Authorization header kullanıldı ve API'nin döndürdüğü JSON cevap frontend tarafına aktarılacak şekilde hazırlandı.

Bu entegrasyonda 401 Unauthorized ve 429 Too Many Requests hataları üzerinde özellikle durdum. 401 hatasının çoğunlukla hatalı token formatı veya eksik API anahtarından, 429 hatasının ise kısa sürede çok fazla istek gönderilmesinden kaynaklanabileceğini gördüm. Rate limit sorununu azaltmak için başarılı cevapların kısa süre cache'lenmesi, isteklerin sıraya alınması ve hata durumunda kullanıcıya anlaşılır mesaj dönülmesi gibi iyileştirmeler yapıldı.

Frontend tarafında DashboardPage ve api.ts üzerinden piyasa verilerinin çekilmesini kontrol ettim. Veriler gelmediğinde boş iskelet ekran yerine hata mesajı gösterilmesi, kullanıcı açısından sorunun daha anlaşılır olmasını sağladı. Aynı gün kart ekranındaki listeleme ve işlem geri bildirimleri de gözden geçirildi. Böylece dış API entegrasyonu, backend servis yapısı ve frontend görüntüleme akışı birlikte test edilmiş oldu.""",
    20: """Yirminci gün projede son kontrolleri yaptım ve final sunumu için hazırlık tamamladım. Uygulamanın genel dilini Türkçe ifadelerle daha tutarlı hale getirdim. Kullanıcıya görünen butonlar, hata mesajları ve yönlendirmeler incelenerek bankacılık uygulamasına daha uygun bir anlatım seçildi. Transfer ekranında "Özete geç" gibi doğal durmayan ifadeler yerine işlemin akışına uygun daha sade metinler kullanıldı.

Para transferi ekranında IBAN doğrulamasını kullanıcı yazarken geri bildirim verecek şekilde iyileştirdim. IBAN'ın TR ile başlaması ve toplam uzunluğunun doğru olması gerektiği anında gösterildi. Böylece kullanıcı hatayı yalnızca işlem sonunda değil, bilgiyi girerken görebildi. Bu küçük değişikliğin kullanıcı deneyimini belirgin biçimde iyileştirdiğini gözlemledim.

Arayüz tarafında uygulamanın renk düzeni kırmızı ve beyaz ağırlıklı olacak şekilde güncellendi. Sol menünün ekran boyunca devam etmesi, kartların daha okunabilir görünmesi ve genel ekran düzeninin bankacılık uygulamasına daha uygun hale gelmesi sağlandı. Dashboard ekranında hesaplar, desteklenen işlemler ve piyasa verileri birlikte kontrol edildi. Dış API verilerinin gelmediği durumlarda bunun kullanıcıya anlaşılır biçimde gösterilmesi üzerinde duruldu.

Sunum hazırlığında proje akışını frontend, controller, service, AppDbContext ve veritabanı sırasıyla anlatacak şekilde düzenledim. Kullanıcı kaydı, giriş, JWT, hesap oluşturma, para yatırma, para çekme, transfer, kart işlemleri, Docker ve dış API entegrasyonu için kısa açıklamalar hazırladım. Sunum sırasında Digital Banking projesinin .NET Web API öğrenme sürecimde controller-service yapısını, veritabanı işlemlerini ve React ile backend iletişimini bir arada görmemi sağladığını anlattım."""
}


DAY_APPENDICES = {
    1: """Bu ilk çalışma sırasında not alırken özellikle "bir değişiklik yapılmadan önce ortamı tanıma" alışkanlığının önemli olduğunu fark ettim. Şirket sistemlerinde yetkim olmadığı için sunucuların hangi amaçla kullanıldığını, bağlantı bilgilerinin nerede tutulduğunu ve işlemlerin kim tarafından onaylandığını gözlem düzeyinde inceledim. Bu nedenle stajın ilk gününü daha çok kurum yapısını, kullanılan araçları ve ekip içindeki iş akışını anlamaya ayırdım. İlerleyen günlerde kendi localimde yapacağım denemeleri ve şirkette takip edeceğim teknik süreçleri daha doğru yorumlayabilmem için bu hazırlık aşaması gerekli bir başlangıç oldu.""",
    2: """RDP bağlantılarını denerken bağlantı kurulamadığında hatanın kaynağını ayırmayı da denedim. Önce sunucunun açık olup olmadığını, sonra IP erişimini, ardından kullanıcı yetkisini ve güvenlik duvarı kuralını kontrol etmek daha düzenli bir yöntem sağladı. Bu yaklaşım bana sistem yönetiminde sorun çözmenin çoğu zaman tahminle değil, olası nedenleri sıraya koyup eleyerek ilerlediğini gösterdi. Günün sonunda bir sunucuya uzaktan erişmenin arkasındaki temel koşulları daha bilinçli yorumlayabildim.""",
    3: """Bu süreçte yerel geliştirme ortamı ile sunucu ortamı arasındaki farkları da karşılaştırdım. Geliştiricinin bilgisayarında çalışan uygulama, sunucuda farklı izinler, farklı portlar veya farklı runtime ayarları nedeniyle hata verebiliyordu. Bu nedenle yayınlama işleminden sonra sadece dosyaların kopyalanması değil, uygulamanın gerçekten cevap verip vermediğinin ve veritabanı gibi bağımlılıklarına erişip erişemediğinin kontrol edilmesi gerekiyordu. Deployment konusunu ilk kez uçtan uca takip etmek, ileride kendi projemi yayınlarken nelere dikkat etmem gerektiğini netleştirdi.""",
    4: """Ağ yapılandırmasıyla ilgili çalışmalarda dokümantasyon tutmanın önemini de gördüm. Hangi IP adresinin hangi sunucuya ait olduğu, hangi portların açıldığı ve yapılan değişikliğin ne zaman uygulandığı kayıt altına alınmadığında hata takibi zorlaşıyordu. Bu yüzden işlem öncesi mevcut ayarları not almak, işlem sonrası da erişim testlerini tekrarlamak gerektiğini öğrendim. Hata giderme sürecindeki bu düzenli kontrol mantığı, yazılım geliştirirken de benzer şekilde log okuma ve adım adım test yapma alışkanlığı kazandırdı.""",
    5: """Uygulama ve veritabanı bağlantısını birlikte değerlendirirken ortam ayrımını da konuştuk. Geliştirme, test ve canlı ortamların farklı ayarlara sahip olabileceğini; özellikle connection string gibi bilgilerin her ortam için ayrı tutulması gerektiğini gördüm. Bu düşünce daha sonra Digital Banking projesindeki appsettings, appsettings.Development.json ve environment variable kullanımını anlamamı kolaylaştırdı. Beşinci gün yaptığım çalışma, uygulamanın yalnızca koddan değil, doğru yapılandırılmış çalışma ortamından da etkilendiğini gösterdi.""",
    6: """Canlı ortama alma sırasında yapılan kontrollerin bir sıra içinde ilerlemesi dikkatimi çekti. Önce ağ erişimi, sonra servislerin durumu, ardından güncelleme ve log izleme adımları kontrol edildi. Bu sıra bozulduğunda sorunun hangi aşamada çıktığını anlamak zorlaşabiliyordu. SQL Server tarafındaki yetki mantığı ile Linux sunuculardaki erişim kontrollerini birlikte düşününce, sistem yönetiminde güvenlik ve sürekliliğin aynı anda gözetildiğini gördüm. Bu deneyim, gerçek ortamda yapılan küçük ayarların bile planlı uygulanması gerektiğini gösterdi.""",
    7: """Docker çalışması sırasında geleneksel kurulumla container yaklaşımı arasındaki bakım farkını da değerlendirdim. Klasik kurulumda veritabanı doğrudan sunucuya bağlıyken, container yapısında servis daha taşınabilir ve izole hale geliyordu. Ancak bu kolaylığın doğru volume tanımı, port yönetimi ve log takibiyle desteklenmesi gerektiğini gördüm. Linux komutlarını kullanarak container ve sunucu durumunu kontrol etmek, komut satırına olan hakimiyetimi artırdı ve sonraki Dockerfile/docker-compose çalışmalarına temel hazırladı.""",
    8: """Proje planlamasında özellikleri tek seferde tamamlamaya çalışmak yerine önceliklendirme yapmanın daha mantıklı olduğunu gördüm. Önce kullanıcı kaydı ve giriş, ardından hesaplar, para hareketleri, kartlar ve arayüz gibi sıralı bir yol izlenmesi kararlaştırıldı. Bu sıralama sayesinde backendde üretilen her endpointin frontendde karşılığını daha sonra kurmak mümkün olacaktı. Yazılım ekibine geçtiğim bu gün, sistem tarafında gördüğüm altyapı bilgisini artık kendi geliştireceğim uygulamaya bağlamaya başladığım gündü.""",
    9: """Mimari tasarım sırasında banka uygulamasındaki verinin hassas yapısını da dikkate aldım. Hesap bakiyesi, IBAN, işlem geçmişi ve kullanıcı bilgileri gibi alanların hem doğru tiplerle tutulması hem de gereksiz şekilde dışarıya açılmaması gerekiyordu. Bu nedenle entity ve DTO kavramlarını daha sonra ayrı ele almak üzere not aldım. İlk taslakta ideal katmanları belirlemek, daha sonra proje kapsamına göre sadeleşme yapılacak olsa bile hangi sorumluluğun nerede durması gerektiğini anlamamı sağladı.""",
    10: """Swagger testleri sırasında endpointlerin dokümantasyon açısından da değerli olduğunu gördüm. Bir controller metodu doğru route ve DTO ile yazıldığında Swagger üzerinde otomatik olarak denenebilir hale geliyordu. Bu durum hem geliştirme sürecinde hızlı test yapmayı hem de projeyi anlatırken hangi endpointin ne iş yaptığını göstermeyi kolaylaştırdı. Kimlik doğrulama akışını planlarken ileride frontend tarafının bu endpointlere nasıl istek göndereceğini de düşünmeye başladım.""",
    11: """Bu gün çalıştığım kavramlar, sonraki kodları okurken doğrudan karşıma çıktı. Örneğin bir controller constructor içinde IAccountService aldığında bunun DI ile sağlandığını, async metotlarda await kullanılmasının veritabanı sorgularını beklemek için gerekli olduğunu daha rahat yorumladım. LINQ metotlarının Türkçe olarak ne işe yaradığını kendi notlarıma yazarak ilerledim. Böylece kodu ezberlemek yerine, her satırın projedeki görevini açıklayabilecek seviyede anlamaya başladım.""",
    12: """Bu akışı örnek bir hesap listeleme isteği üzerinden düşündüm. Frontend istek gönderiyor, AccountsController isteği karşılıyor, token içinden kullanıcı kimliği alınıyor, IAccountService üzerinden AccountService metodu çalışıyor ve AppDbContext ile veritabanından hesaplar okunuyordu. Sonuç AccountResponseDto listesi halinde istemciye dönüyordu. Bu basit zinciri kurmak, sonraki sunumda "frontend backend ile nasıl konuşuyor" sorusunu açıklayabilmem için temel bir anlatım haline geldi.""",
    13: """Refactoring çalışması sırasında dosya adlarının ve klasörlerin anlaşılır olmasının ekip çalışmasındaki etkisini de fark ettim. Bir işlemle ilgili controller, service, interface, DTO ve model aynı isim mantığıyla ilerlediğinde kodda gezinmek daha kolay oluyordu. Örneğin transfer işlemi için TransactionsController, ITransactionService, TransactionService, TransferRequestDto ve Transaction entity'si arasında açık bir bağ kurulabiliyordu. Bu düzen, projeyi sunarken "hangi işlem nerede yapılıyor" sorusuna daha rahat cevap verebilmemi sağladı.""",
    14: """Yetkilendirme testlerinde frontend veya Swagger üzerinden token gönderilmediğinde korumalı endpointlerin çalışmadığını gördüm. Bu davranış başlangıçta hata gibi görünse de aslında güvenlik için beklenen sonuçtu. Daha önce dış API tarafında yaşanan 401 hatasını da bu bilgiyle daha kolay yorumlayabildim; 401 cevabı çoğu zaman kimlik doğrulama bilgisinin eksik ya da hatalı olduğunu gösteriyordu. JWT konusunu bu şekilde uygulama üzerinde görmek, teorik açıklamadan daha kalıcı oldu.""",
    15: """Transfer işlemini geliştirirken veritabanı tutarlılığı konusu özellikle önemliydi. Para bir hesaptan çıkıp diğerine girerken işlemin yarım kalmaması gerekiyordu. Bu nedenle bakiye güncellemeleri ve transaction kaydı tek bir akış olarak ele alındı. Hatalı tutar, yetersiz bakiye, bulunmayan IBAN veya aynı hesaba transfer gibi durumlarda kullanıcıya açıklayıcı mesaj dönülmesi planlandı.""",
    16: """Frontend geliştirmesinde tekrar eden arayüz parçalarını component olarak kullanmanın düzen sağladığını gördüm. Button, Input, Card, Modal, Table ve EmptyState gibi ortak bileşenler sayesinde sayfalar arasında daha tutarlı bir görünüm elde edildi. API çağrılarını tek dosyada toplamak da bakım açısından kolaylık sağladı; endpoint adresi veya token gönderme şekli değiştiğinde tüm sayfaları ayrı ayrı düzeltmek gerekmiyordu. Bu yapı, React tarafındaki kodun backendle daha kontrollü iletişim kurmasına yardımcı oldu.""",
    17: """Docker tarafında environment variable kullanımının yalnızca kolaylık değil güvenlik açısından da önemli olduğunu değerlendirdim. Veritabanı bağlantısı, JWT secret ve dış API anahtarı gibi değerlerin imajın içine sabitlenmesi doğru değildi. docker-compose üzerinden bu bilgilerin çalışma zamanında verilmesi, aynı imajın farklı ortamlarda farklı ayarlarla kullanılmasını sağlıyordu. Bu gün yapılan düzenlemeler, projenin sadece yerel bilgisayarda değil daha taşınabilir bir ortamda da çalıştırılabileceğini gösterdi.""",
    18: """Migration dosyalarını incelerken veritabanı şemasının zaman içindeki değişimini takip edebilmenin avantajını gördüm. İlk migration temel tabloları oluştururken sonraki migrationlarda kart işlemleri ve tablo adlandırmaları gibi değişiklikler yer aldı. Bu dosyalar sayesinde veritabanı yapısının hangi aşamada nasıl değiştiği anlaşılabiliyordu. AppDbContext ve migration ilişkisini kavramak, "veritabanı işlemleri nerede yapılıyor" sorusunu yalnızca tablo üzerinden değil, kod ve migration süreci üzerinden açıklayabilmemi sağladı.""",
    19: """Dış API entegrasyonunda frontendin doğrudan CollectAPI'ye gitmemesinin daha doğru olduğunu gördüm. İstek önce kendi backendimizdeki /api/market endpointlerine geliyor, backend MarketService içinde dış servise bağlanıyor ve sonucu frontend için uygun biçimde döndürüyordu. Bu yapı API anahtarının tarayıcı tarafında görünmesini engelliyor ve hata yönetimini backendde toplamayı sağlıyordu. Market verilerinin gelmediği durumlarda cache, timeout ve kullanıcı mesajı gibi kontrollerin uygulamanın daha kararlı görünmesine katkı sağladığını değerlendirdim.""",
    20: """Final sunumunda çok ileri kod ayrıntılarına girmeden projenin katmanlarını ve akışlarını anlatmaya odaklandım. Her özellik için "amaç nedir, endpoint nerede, hangi service çalışır, DTO ve entity hangisidir, veritabanı işlemi nerede yapılır" sorularına kısa cevaplar hazırladım. Projenin geliştirilebilecek yönleri olarak daha kapsamlı testler, daha güvenli secret yönetimi, daha detaylı validasyonlar ve dış API için daha güçlü hata yönetimi not edildi. Böylece stajın son günü hem teknik kontrol hem de anlatım hazırlığıyla tamamlandı."""
}


TOPIC_OVERRIDES = {
    3: "ASP.NET Uygulamasının Local Ortamda Yayınlanması",
    5: "ASP.NET Uygulama Geliştirme, Local Yayınlama ve Veritabanı Süreci",
}


def clear_cell(cell):
    for paragraph in list(cell.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)


def format_paragraph(paragraph):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def replace_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_body(cell, text):
    clear_cell(cell)
    parts = [part.strip() for part in text.split("\n\n") if part.strip()]
    for index, part in enumerate(parts):
        paragraph = cell.add_paragraph()
        if index:
            paragraph.paragraph_format.space_before = Pt(4)
        run = paragraph.add_run(part)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_after = Pt(0)


def apply_daily_formatting(doc):
    for table_index, day in enumerate(range(1, 21), start=15):
        table = doc.tables[table_index]
        if day in TOPIC_OVERRIDES:
            replace_cell_text(table.rows[1].cells[1], TOPIC_OVERRIDES[day])
        body_cell = table.rows[3].cells[0]
        add_body(body_cell, f"{DAY_BODIES[day]}\n\n{DAY_APPENDICES[day]}")


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    apply_daily_formatting(doc)
    try:
        doc.save(OUTPUT)
        print(OUTPUT)
    except PermissionError:
        doc.save(FALLBACK_OUTPUT)
        print(FALLBACK_OUTPUT)


if __name__ == "__main__":
    main()
