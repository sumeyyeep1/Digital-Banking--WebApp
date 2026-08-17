import re
from pathlib import Path

import pdfplumber
from docx import Document
from PIL import Image, ImageDraw, ImageFont


DOCX = Path(r"C:\Users\sümeyye\Desktop\DigitalBankingSimple\outputs\staj_defteri_duzenlenmis.docx")
PDF = Path(r"C:\Users\sümeyye\Desktop\DigitalBankingSimple\outputs\staj_defteri_duzenlenmis.pdf")
PAGES_DIR = Path(r"C:\Users\sümeyye\Desktop\DigitalBankingSimple\outputs\staj_pages")
SHEETS_DIR = Path(r"C:\Users\sümeyye\Desktop\DigitalBankingSimple\outputs\staj_contact_sheets")


def inspect_docx():
    doc = Document(DOCX)
    print(f"docx_tables={len(doc.tables)}")
    problems = []
    for table_index, day in enumerate(range(1, 21), start=15):
        table = doc.tables[table_index]
        cell = table.rows[3].cells[0]
        text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
        if len(text) < 1200:
            problems.append(f"day {day}: body too short ({len(text)})")
        for p in cell.paragraphs:
            if p.text.strip() and p.paragraph_format.line_spacing_rule is None:
                problems.append(f"day {day}: missing explicit line spacing")
            for run in p.runs:
                if run.text.strip() and run.font.size and run.font.size.pt != 12:
                    problems.append(f"day {day}: run font size {run.font.size.pt}")
    print("docx_daily_tables=20")
    print("docx_problems=" + ("; ".join(problems) if problems else "none"))


def inspect_pdf():
    with pdfplumber.open(PDF) as pdf:
        print(f"pdf_pages={len(pdf.pages)}")
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            dates = re.findall(r"\b\d{2}\.\d{2}\.\d{4}\b", text)
            day_labels = re.findall(r"\b\d{1,2}\.\s*GÜN\b", text, flags=re.IGNORECASE)
            if dates or day_labels:
                print(f"page {index:02d}: dates={','.join(dates) or '-'} labels={','.join(day_labels) or '-'}")


def make_contact_sheets():
    SHEETS_DIR.mkdir(parents=True, exist_ok=True)
    images = sorted(PAGES_DIR.glob("page-*.png"), key=lambda p: int(re.search(r"-(\d+)\.png$", p.name).group(1)))
    if not images:
        print("contact_sheets=none")
        return
    thumb_w = 310
    gap = 18
    header_h = 36
    cols = 3
    font = ImageFont.load_default()
    sheet_paths = []
    for chunk_start in range(0, len(images), 9):
        chunk = images[chunk_start : chunk_start + 9]
        thumbs = []
        for path in chunk:
            img = Image.open(path).convert("RGB")
            scale = thumb_w / img.width
            thumb = img.resize((thumb_w, int(img.height * scale)))
            thumbs.append((path, thumb))
        rows = (len(thumbs) + cols - 1) // cols
        cell_h = max(t.height for _, t in thumbs) + header_h
        sheet = Image.new("RGB", (cols * thumb_w + (cols + 1) * gap, rows * cell_h + (rows + 1) * gap), "white")
        draw = ImageDraw.Draw(sheet)
        for i, (path, thumb) in enumerate(thumbs):
            row, col = divmod(i, cols)
            x = gap + col * (thumb_w + gap)
            y = gap + row * (cell_h + gap)
            page_no = int(re.search(r"-(\d+)\.png$", path.name).group(1))
            draw.text((x, y), f"Sayfa {page_no}", fill="black", font=font)
            sheet.paste(thumb, (x, y + header_h))
        out = SHEETS_DIR / f"contact-{chunk_start // 9 + 1}.png"
        sheet.save(out)
        sheet_paths.append(str(out))
    print("contact_sheets=" + " | ".join(sheet_paths))


if __name__ == "__main__":
    inspect_docx()
    inspect_pdf()
    make_contact_sheets()
