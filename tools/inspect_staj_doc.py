from docx import Document
from pathlib import Path

path = Path(r"C:\Users\sümeyye\Downloads\klu_staj_defteri_şablon_-__temmuz_2026__ (1).docx")
doc = Document(path)

print("paragraphs", len(doc.paragraphs), "tables", len(doc.tables), "sections", len(doc.sections))

for ti, table in enumerate(doc.tables):
    print(f"\nTABLE {ti} rows {len(table.rows)} cols {len(table.columns)}")
    for ri, row in enumerate(table.rows):
        vals = []
        for ci, cell in enumerate(row.cells):
            text = " / ".join(x.strip() for x in cell.text.splitlines() if x.strip())
            if len(text) > 220:
                text = text[:220] + "..."
            vals.append(f"c{ci}={text!r}")
        print(" row", ri, " | ".join(vals))
        if ri > 12:
            print(" ...")
            break
